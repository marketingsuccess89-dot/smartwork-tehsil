import io
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

FONT_NAME = 'Nirmala UI'  # Industry standard Devanagari Hindi font in Windows & Office

def apply_font(run, font_name=FONT_NAME):
    """
    Applies font name across Latin, Complex Script (Devanagari/Hindi), and East Asian
    to guarantee Microsoft Word renders Hindi in Nirmala UI instead of falling back to Mangal.
    """
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = rpr._add_rFonts()
    rfonts.set(qn('w:ascii'), font_name)
    rfonts.set(qn('w:hAnsi'), font_name)
    rfonts.set(qn('w:cs'), font_name)
    rfonts.set(qn('w:eastAsia'), font_name)

def set_cell_background(cell, fill_hex="F3F4F6"):
    """Sets background shading of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_table_styling(table, is_signature=False):
    """Applies clean legal document styling to tables."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = table._tbl.tblPr
    if is_signature:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="none"/>
                <w:left w:val="none"/>
                <w:bottom w:val="none"/>
                <w:right w:val="none"/>
                <w:insideH w:val="none"/>
                <w:insideV w:val="none"/>
            </w:tblBorders>
        ''')
    else:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="6" w:space="0" w:color="06281E"/>
                <w:left w:val="none"/>
                <w:bottom w:val="single" w:sz="6" w:space="0" w:color="06281E"/>
                <w:right w:val="none"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>
                <w:insideV w:val="none"/>
            </w:tblBorders>
        ''')
    tblPr.append(borders)

def add_styled_paragraph(doc, text: str, style_type='body', alignment=None):
    """Adds a paragraph parsing inline **bold** markdown tags."""
    p = doc.add_paragraph()
    
    if alignment is not None:
        p.alignment = alignment
    elif style_type == 'body':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p_format = p.paragraph_format
    if style_type == 'title':
        p_format.space_before = Pt(12)
        p_format.space_after = Pt(14)
        p_format.line_spacing = 1.2
    elif style_type == 'heading':
        p_format.space_before = Pt(14)
        p_format.space_after = Pt(6)
        p_format.line_spacing = 1.15
    elif style_type == 'subheading':
        p_format.space_before = Pt(8)
        p_format.space_after = Pt(3)
        p_format.line_spacing = 1.15
    elif style_type == 'clause':
        p_format.space_before = Pt(4)
        p_format.space_after = Pt(6)
        p_format.line_spacing = 1.15
    else:
        p_format.space_before = Pt(2)
        p_format.space_after = Pt(6)
        p_format.line_spacing = 1.15

    # Clean up <br> tags into newlines for Word formatting
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)

    # Parse **bold** text
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        is_bold = False
        content = part
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            is_bold = True
            content = part[2:-2]

        run = p.add_run(content)
        apply_font(run)
        run.bold = is_bold

        if style_type == 'title':
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(6, 40, 30) # Rich Deep Emerald
        elif style_type == 'heading':
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(10, 25, 20)
        elif style_type == 'subheading':
            run.font.size = Pt(13)
            run.font.bold = True
        else:
            run.font.size = Pt(12)

    return p

def render_markdown_table(doc, table_lines):
    """Renders a parsed markdown table into a beautifully formatted docx Table."""
    raw_rows = []
    for l in table_lines:
        s = l.strip()
        if not s or re.match(r'^[\|\s\-:]+$', s):
            continue
        cells = [c.strip() for c in s.strip('|').split('|')]
        raw_rows.append(cells)

    if not raw_rows:
        return

    num_cols = max(len(r) for r in raw_rows)
    num_rows = len(raw_rows)

    # If this is a dummy 2-column table used only to push a single person's details to the right
    # (i.e. column 0 is completely empty across all rows), output clean right-aligned paragraphs
    # rather than creating an unwanted table in MS Word!
    if num_cols == 2 and all((len(r) < 1 or not r[0].strip()) for r in raw_rows):
        for r in raw_rows:
            txt = r[1].strip() if len(r) > 1 else ""
            if txt:
                add_styled_paragraph(doc, txt, style_type='body', alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        return

    sig_keywords = (
        'हस्ताक्षर', 'हसताक्षर', 'हस्तक्षर', 'हस्ताक्षरी', 'साक्षी', 'साक्क्षी', 'गवाह',
        'प्रथम पक्ष', 'परथम पक्ष', 'द्वितीय पक्ष', 'दवतीय पक्ष', 'क्रेता', 'विक्रेता',
        'शपथकर्ता', 'आवेदक', 'प्रार्थी', 'मुवक्किल', 'अधिवक्ता', 'निवेदक', 'भवदीय',
        'signature', 'witness', 'first party', 'second party', 'landlord', 'tenant',
        'buyer', 'seller', 'deponent', 'applicant', 'sincerely', 'regards'
    )
    is_sig_table = any(any(kw in cell.lower() for kw in sig_keywords) for row in raw_rows for cell in row)
    # If 2 columns and short (e.g. party names/signatures), default to borderless
    if num_cols == 2 and not is_sig_table and num_rows <= 4:
        is_sig_table = True

    table = doc.add_table(rows=num_rows, cols=num_cols)
    set_table_styling(table, is_signature=is_sig_table)

    for r_idx, row_data in enumerate(raw_rows):
        row = table.rows[r_idx]
        is_header = (r_idx == 0 and not is_sig_table)

        for c_idx in range(num_cols):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell_text = row_data[c_idx] if c_idx < len(row_data) else ""

            if is_header:
                set_cell_background(cell, "F0FDF4") # Subtle emerald tint header

            # Format cell paragraphs
            p = cell.paragraphs[0]
            if is_sig_table and num_cols == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
            elif is_header or is_sig_table:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15

            # Handle <br> tags within cells
            sub_lines = cell_text.replace('<br>', '\n').replace('<br/>', '\n').split('\n')
            for s_idx, s_line in enumerate(sub_lines):
                if s_idx > 0:
                    p = cell.add_paragraph()
                    if is_sig_table and num_cols == 2:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx == 0 else WD_ALIGN_PARAGRAPH.RIGHT
                    elif is_sig_table:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.line_spacing = 1.15

                parts = re.split(r'(\*\*.*?\*\*)', s_line)
                for part in parts:
                    if not part:
                        continue
                    bold = is_header
                    txt = part
                    if part.startswith('**') and part.endswith('**') and len(part) >= 4:
                        bold = True
                        txt = part[2:-2]
                    
                    run = p.add_run(txt)
                    apply_font(run)
                    run.font.size = Pt(10.5)
                    run.bold = bold

def unwrap_paragraphs(text: str) -> str:
    """
    Combines artificial mid-sentence line breaks from handwritten notes
    into full, continuous flowing A4 paragraphs without losing document structure.
    Universal for all documents: Letters, Applications, Deeds, Affidavits.
    """
    lines = text.split('\n')
    unwrapped = []
    current_p = []
    in_closing = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current_p:
                unwrapped.append(' '.join(current_p))
                current_p = []
            unwrapped.append('')
            continue

        is_heading = stripped.startswith('#')
        is_bullet = stripped.startswith('* ') or stripped.startswith('- ')
        is_table = stripped.startswith('|')
        is_page_break = bool(re.match(r'^\s*-{3,}\s*$', stripped))
        if is_page_break:
            in_closing = False

        clean_stripped = re.sub(r'[*#_]', '', stripped).strip()

        is_new_clause = bool(re.match(r'^(?:(?:\(?(\d+|[०-९]+|[क-ह])\))|(\d+|[०-९]+)[\.\)])\s+', clean_stripped))
        # Universal label line: e.g. "नाम :", "विषय :", "पता :", "कक्षा :", "दिनांक :", "आवेदक / प्रार्थी:"
        is_label_line = bool(re.match(r'^[^:\n]{2,35}\s*:\s*.*$', clean_stripped))

        # Precision closing detection: sentences ending in verbs are NOT closing blocks
        is_sentence = bool(re.search(r'(?:कि:|है[।\.]|हूँ[।\.]|था[।\.]|करें[।\.]|गया[।\.]|जाएगा[।\.])$', clean_stripped))
        is_closing_start = False
        if not is_sentence and len(clean_stripped) < 45:
            if re.match(r'^(?:द्वारा अधिवक्ता|अधिवक्ता|हस्ताक्षर|भवदीय|निवेदक|शपथी|शपथकर्ता|विनीत|आपका आज्ञाकारी|आज्ञाकारी|स्वीकृत व प्रस्तुतकर्ता|Sincerely|Regards|Yours obediently|Yours faithfully)\b', clean_stripped, re.IGNORECASE):
                is_closing_start = True
            elif re.match(r'^(?:आवेदक|प्रार्थी)\s*(?:[/:,।\-]|बनाम|$)', clean_stripped, re.IGNORECASE) and not re.search(r'(?:सादर|निवेदन|प्रार्थना|करता|करती)', clean_stripped):
                is_closing_start = True

        if is_closing_start:
            in_closing = True
        elif in_closing and (is_sentence or len(clean_stripped) > 60 or is_new_clause or is_heading):
            in_closing = False

        # Universal formal opening / court headers (excluding body sentences starting with applicant name)
        is_formal_break = bool(re.match(r'^(सेवा में|महोदय|महोदया|श्रीमान|मान्यवर|विषय|स्थान|दिनांक|न्यायालय|मुकदमा|बनाम|थाना|धारा|प्रार्थना|अनुतोष|स्वीकृत|संलग्नक|नाम|पिता|पता|कक्षा|अनुक्रमांक|मो०|मोबाइल|To:|From:|Subject:|Dear|Respected|Date:)', clean_stripped, re.IGNORECASE)) or is_closing_start

        if in_closing or is_heading or is_bullet or is_table or is_page_break or is_new_clause or is_label_line or is_formal_break:
            if current_p:
                unwrapped.append(' '.join(current_p))
                current_p = []
            unwrapped.append(stripped)
        else:
            if current_p and not (current_p[-1].startswith('#') or current_p[-1].startswith('|') or re.match(r'^\s*-{3,}\s*$', current_p[-1]) or re.match(r'^[^:\n]{2,35}\s*:\s*.*$', current_p[-1]) or re.match(r'^(सेवा में|महोदय|महोदया|श्रीमान|मान्यवर|विषय|भवदीय|प्रार्थी|आवेदक|निवेदक|शपथी|हस्ताक्षर|न्यायालय|मुकदमा|बनाम|To:|From:|Subject:|Dear|Respected|Sincerely)', current_p[-1], re.IGNORECASE)):
                current_p.append(stripped)
            else:
                if current_p:
                    unwrapped.append(' '.join(current_p))
                    current_p = []
                current_p.append(stripped)

    if current_p:
        unwrapped.append(' '.join(current_p))

    return '\n'.join(unwrapped)

def create_docx(text: str, stamp_paper: bool = False) -> io.BytesIO:
    """
    Converts legal draft markdown into a styled, professional MS Word (.docx) document
    ready for Indian Tehsil / Court registry printing.
    """
    doc = Document()

    # Configure first page margins (Standard Normal Margins: 1.0 in / 2.54 cm)
    section = doc.sections[0]
    section.top_margin = Inches(3.0 if stamp_paper else 1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    text = unwrap_paragraphs(text)
    lines = text.split('\n')
    i = 0
    in_closing_block = False
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 1. Section break / Multi-page stamp layout (---)
        if re.match(r'^\s*-{3,}\s*$', stripped):
            new_section = doc.add_section()
            new_section.top_margin = Inches(1.0)
            new_section.bottom_margin = Inches(1.0)
            new_section.left_margin = Inches(1.0)
            new_section.right_margin = Inches(1.0)
            in_closing_block = False
            i += 1
            continue

        # 2. Empty line
        if not stripped:
            i += 1
            continue

        # 3. Markdown Tables (e.g. Signatures, Witnesses, Consideration breakdown)
        if stripped.startswith('|') and stripped.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            render_markdown_table(doc, table_lines)
            continue

        # 4. Heading 1 (Title) e.g. # Title
        if stripped.startswith('# '):
            in_closing_block = False
            add_styled_paragraph(doc, stripped[2:], style_type='title', alignment=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue

        # 5. Heading 2 e.g. ## Section
        if stripped.startswith('## '):
            in_closing_block = False
            add_styled_paragraph(doc, stripped[3:], style_type='heading', alignment=WD_ALIGN_PARAGRAPH.LEFT)
            i += 1
            continue

        # 6. Heading 3 e.g. ### Sub-section
        if stripped.startswith('### '):
            add_styled_paragraph(doc, stripped[4:], style_type='subheading', alignment=WD_ALIGN_PARAGRAPH.LEFT)
            i += 1
            continue

        # 7. Bullet list items (* or -)
        if stripped.startswith('* ') or stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            content = stripped[2:]
            parts = re.split(r'(\*\*.*?\*\*)', content)
            for part in parts:
                if not part: continue
                is_bold = part.startswith('**') and part.endswith('**') and len(part) >= 4
                txt = part[2:-2] if is_bold else part
                run = p.add_run(txt)
                apply_font(run)
                run.font.size = Pt(12)
                run.bold = is_bold
            i += 1
            continue

        # 8. Numbered list items (e.g. 1., १., (1), (१), (क))
        match = re.match(r'^(?:(?:\(?(\d+|[०-९]+|[क-ह])\))|(\d+|[०-९]+)[\.\)])\s+(.*)$', stripped)
        if match:
            num = match.group(1) or match.group(2)
            content = match.group(3)
            add_styled_paragraph(doc, f"**{num}.** {content}", style_type='clause', alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)
            i += 1
            continue

        # 9. Smart universal detection for multiple side-by-side signatures on a single line
        if stripped.count('हस्ताक्षर') >= 2 and len(stripped) > 15:
            next_l1 = lines[i+1].strip() if i + 1 < len(lines) else ""
            next_l2 = lines[i+2].strip() if i + 2 < len(lines) else ""

            parts = re.split(r'\s{3,}|\t|(?=द्वितीय|मकान मालिक|क्रेता|पक्षकार)', stripped)
            parts = [p.strip() for p in parts if p.strip()]
            h_left = parts[0] if len(parts) > 0 else "हस्ताक्षर"
            h_right = parts[1] if len(parts) > 1 else "हस्ताक्षर"

            n_left, n_right = "", ""
            consumed_l1 = False
            if next_l1 and ('हस्ताक्षर' not in next_l1) and ('दिनांक' not in next_l1):
                n_parts = re.split(r'\s{2,}|\t', next_l1)
                n_parts = [p.strip() for p in n_parts if p.strip()]
                if len(n_parts) >= 2:
                    n_left = n_parts[0]
                    n_right = n_parts[1]
                elif len(n_parts) == 1:
                    n_left = n_parts[0]
                consumed_l1 = True

            d_left, d_right = "", ""
            consumed_l2 = False
            if consumed_l1 and next_l2 and 'दिनांक' in next_l2:
                d_parts = re.split(r'(?=दिनांक)', next_l2)
                d_parts = [p.strip() for p in d_parts if p.strip()]
                if len(d_parts) >= 2:
                    d_left = d_parts[0]
                    d_right = d_parts[1]
                elif len(d_parts) == 1:
                    d_left = d_parts[0]
                consumed_l2 = True

            left_cell = f"**{h_left}**<br><br>___________________"
            if n_left: left_cell += f"<br>{n_left}"
            if d_left: left_cell += f"<br>{d_left}"

            right_cell = f"**{h_right}**<br><br>___________________"
            if n_right: right_cell += f"<br>{n_right}"
            if d_right: right_cell += f"<br>{d_right}"

            sig_table_lines = [
                f"| {h_left} | {h_right} |",
                "| :--- | ---: |",
                f"| {left_cell} | {right_cell} |"
            ]
            render_markdown_table(doc, sig_table_lines)
            i += 1
            if consumed_l1: i += 1
            if consumed_l2: i += 1
            continue

        # 10. Single closing / signature / applicant block detection (e.g. भवदीय, प्रार्थी, शपथकर्ता)
        clean_stripped = re.sub(r'[*#_]', '', stripped).strip()
        if re.match(r'^(दिनांक|स्थान|Date:|Place:)', clean_stripped, re.IGNORECASE):
            add_styled_paragraph(doc, stripped, style_type='body', alignment=WD_ALIGN_PARAGRAPH.LEFT)
            i += 1
            continue

        is_sentence = bool(re.search(r'(?:कि:|है[।\.]|हूँ[।\.]|था[।\.]|करें[।\.]|गया[।\.]|जाएगा[।\.])$', clean_stripped))
        is_closing_start = False
        if not is_sentence and len(clean_stripped) < 45:
            if re.match(r'^(?:द्वारा अधिवक्ता|अधिवक्ता|हस्ताक्षर|भवदीय|निवेदक|शपथी|शपथकर्ता|विनीत|आपका आज्ञाकारी|आज्ञाकारी|स्वीकृत व प्रस्तुतकर्ता|Sincerely|Regards|Yours obediently|Yours faithfully|Yours truly)\b', clean_stripped, re.IGNORECASE):
                is_closing_start = True
            elif re.match(r'^(?:आवेदक|प्रार्थी)\s*(?:[/:,।\-]|बनाम|$)', clean_stripped, re.IGNORECASE) and not re.search(r'(?:सादर|निवेदन|प्रार्थना|करता|करती)', clean_stripped):
                is_closing_start = True

        if is_closing_start:
            in_closing_block = True
        elif in_closing_block and (is_sentence or len(clean_stripped) > 60 or re.match(r'^(?:(?:\(?(\d+|[०-९]+|[क-ह])\))|(\d+|[०-९]+)[\.\)])\s+', clean_stripped) or stripped.startswith('#')):
            in_closing_block = False

        align = WD_ALIGN_PARAGRAPH.RIGHT if in_closing_block else WD_ALIGN_PARAGRAPH.JUSTIFY
        add_styled_paragraph(doc, stripped, style_type='body', alignment=align)
        i += 1

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
